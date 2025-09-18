import os
import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

TITLE_COLOR = RGBColor(0, 32, 91)       # #00205b
BODY_COLOR = RGBColor(0, 0, 0)          # Black body text
SUBSUBHEADING_COLOR = BODY_COLOR        # Black for sub-subheadings
FOOTER_COLOR = RGBColor(0xE4, 0x00, 0x2B)  # #e4002b red

def generate_docx(body, output_path, title="Generated Document"):
    doc = Document()

    # Clean title of hashes and unwanted chars and print centered once
    clean_title = re.sub(r'^[#*\-_=\s]+|[#*\-_=\s]+$', '', title).strip()
    if clean_title:
        para = doc.add_paragraph()
        run = para.add_run(clean_title)
        run.font.size = Pt(16)
        run.font.bold = True
        run.font.color.rgb = TITLE_COLOR
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.space_after = Pt(12)

    body_clean = clean_text(body)
    sections = parse_sections(body_clean)

    for sec in sections:
        level = sec['level']
        heading = sec['heading']
        content = sec['body']

        if level == 1:
            add_heading(doc, heading, font_size=16, font_color=TITLE_COLOR)
        elif level == 2:
            add_heading(doc, heading, font_size=14, font_color=TITLE_COLOR)
        elif level == 3:
            add_heading(doc, heading, font_size=12, font_color=SUBSUBHEADING_COLOR)

        add_body_content(doc, content)

    add_footer(doc)

    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    print(f"📝 DOCX saved to: {output_path}")

def clean_text(text):
    lines = text.splitlines()
    cleaned_lines = []
    in_table = False
    for line in lines:
        stripped = line.strip()

        # Preserve tables intact
        if stripped.startswith('|') and stripped.endswith('|'):
            in_table = True
            cleaned_lines.append(line)
            continue
        elif in_table and stripped == '':
            in_table = False
            cleaned_lines.append(line)
            continue
        elif in_table:
            cleaned_lines.append(line)
            continue

        # Keep heading hashes (#, ##, ###) only at line start, cleaning their content
        heading_match = re.match(r'^(#{1,3})\s+(.*)$', line)
        if heading_match:
            hashes = heading_match.group(1)
            content = heading_match.group(2)
            content = re.sub(r'[*=_\-\s]+$', '', content)  # strip trailing decorations
            content = re.sub(r'[*_]+', '', content)       # remove * and _ inside heading text
            cleaned_lines.append(f"{hashes} {content.strip()}")
            continue

        # Remove lines of repeated markdown chars used as separators
        if re.match(r'^[=\-_*]{3,}$', stripped):
            continue

        # Remove stray * and _ in other lines
        clean_line = re.sub(r'[*_]+', '', line)
        cleaned_lines.append(clean_line)

    return "\n".join(cleaned_lines)

def parse_sections(text):
    pattern = re.compile(r'^(#{1,3})\s+(.*)$', re.MULTILINE)
    matches = list(pattern.finditer(text))
    if not matches:
        return [{'level': 0, 'heading': '', 'body': text.strip()}]

    sections = []

    # Text before first heading
    if matches[0].start() > 0:
        pre_content = text[:matches[0].start()].strip()
        if pre_content:
            sections.append({'level': 0, 'heading': '', 'body': pre_content})

    for i, match in enumerate(matches):
        level = len(match.group(1))
        heading = match.group(2).strip()
        start = match.end()
        end = matches[i + 1].start() if i + 1 < len(matches) else len(text)
        body = text[start:end].strip()
        sections.append({'level': level, 'heading': heading, 'body': body})

    return sections

def add_heading(doc, text, font_size, font_color):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.size = Pt(font_size)
    run.font.bold = True
    run.font.color.rgb = font_color
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    para.space_after = Pt({16:12, 14:8, 12:6}[font_size])

def add_body_content(doc, content):
    for line in content.splitlines():
        l = line.strip()
        if not l:
            continue
        if re.match(r'^[-+•]\s+', l):
            para = doc.add_paragraph(style='List Bullet')
            run = para.add_run(re.sub(r'^[-+•]\s+', '', l))
            run.font.size = Pt(12)
            run.font.color.rgb = BODY_COLOR
            para.space_after = Pt(4)
        else:
            para = doc.add_paragraph()
            run = para.add_run(l)
            run.font.size = Pt(12)
            run.font.color.rgb = BODY_COLOR
            para.space_after = Pt(6)
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT

def add_footer(doc):
    section = doc.sections[0]
    footer = section.footer
    para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()

    while para.runs:
        para.runs[0].clear()

    run_left = para.add_run("HPGPT")
    run_left.font.size = Pt(10)
    run_left.font.bold = True
    run_left.font.color.rgb = FOOTER_COLOR

    para.add_run("\t")

    run_page = para.add_run()
    fldChar_begin = OxmlElement('w:fldChar')
    fldChar_begin.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"
    fldChar_separate = OxmlElement('w:fldChar')
    fldChar_separate.set(qn('w:fldCharType'), 'separate')
    fldChar_end = OxmlElement('w:fldChar')
    fldChar_end.set(qn('w:fldCharType'), 'end')

    run_page._r.append(fldChar_begin)
    run_page._r.append(instrText)
    run_page._r.append(fldChar_separate)
    run_page._r.append(fldChar_end)

    run_page.font.size = Pt(10)
    run_page.font.bold = True
    run_page.font.color.rgb = FOOTER_COLOR

    para.paragraph_format.tab_stops.clear_all()
    para.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), WD_ALIGN_PARAGRAPH.RIGHT)
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
